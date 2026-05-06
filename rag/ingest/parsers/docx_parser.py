from __future__ import annotations

from .base import BaseParser, ParseResult, TextChunk


class DocxParser(BaseParser):

    def can_parse(self, file_path: str) -> bool:
        return self._ext(file_path) == ".docx"

    def parse(self, file_path: str) -> ParseResult:
        try:
            from docx import Document
        except ImportError:
            return ParseResult(
                source_path=file_path, doc_type="docx",
                errors=["python-docx가 설치되지 않았습니다."]
            )

        chunks: list[TextChunk] = []
        errors: list[str] = []
        title = ""
        para_count = 0
        table_count = 0

        try:
            doc = Document(file_path)

            try:
                title = doc.core_properties.title or ""
            except Exception:
                title = ""

            current_section: list[str] = []
            current_heading = ""
            chunk_idx = 0

            def flush_section():
                nonlocal chunk_idx, current_section, current_heading
                if not current_section:
                    return
                text = "\n".join(current_section).strip()
                if text:
                    header = f"[제목: {current_heading}]\n" if current_heading else ""
                    chunks.append(TextChunk(
                        content=header + text,
                        chunk_index=chunk_idx,
                        metadata={"heading": current_heading},
                    ))
                    chunk_idx += 1
                current_section = []

            for para in doc.paragraphs:
                txt = para.text.strip()
                if not txt:
                    continue
                para_count += 1
                style = (para.style.name or "").lower() if para.style else ""
                if style.startswith("heading"):
                    flush_section()
                    current_heading = txt
                else:
                    current_section.append(txt)
            flush_section()

            for tbl_idx, table in enumerate(doc.tables):
                table_count += 1
                rows_md: list[str] = []
                for r_idx, row in enumerate(table.rows):
                    cells = [cell.text.strip().replace("\n", " ") for cell in row.cells]
                    rows_md.append("| " + " | ".join(cells) + " |")
                    if r_idx == 0:
                        rows_md.append("| " + " | ".join(["---"] * len(cells)) + " |")
                if rows_md:
                    chunks.append(TextChunk(
                        content=f"[표 {tbl_idx + 1}]\n" + "\n".join(rows_md),
                        chunk_index=chunk_idx,
                        metadata={"table_index": tbl_idx},
                    ))
                    chunk_idx += 1
        except Exception as e:
            errors.append(str(e))

        return ParseResult(
            source_path=file_path,
            doc_type="docx",
            chunks=chunks,
            metadata={
                "title": title,
                "paragraph_count": para_count,
                "table_count": table_count,
            },
            errors=errors,
        )
